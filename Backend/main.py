from fastapi import FastAPI, Depends, HTTPException, UploadFile, File
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from jose import JWTError, jwt
from passlib.context import CryptContext
from datetime import datetime, timedelta
import os, json, uuid, tempfile
from cryptography.fernet import Fernet, InvalidToken
import chromadb
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_ollama import OllamaLLM
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
import easyocr
import whisper
import logging
from fastapi.responses import FileResponse, StreamingResponse
from typing import List, Optional
from PIL import Image
import fitz
from io import BytesIO
import asyncio  # Added for background tasks

app = FastAPI()

# --- CORS Configuration ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],  # Add common frontend origins
    allow_credentials=True,
    allow_methods=["*"],  # Allow all methods for flexibility
    allow_headers=["*"],  # Allow all headers for flexibility
)

# --- Security Configuration ---
# !!! IMPORTANT: Replace "your-secret-key" with a strong, random string !!!
SECRET_KEY = "your-secret-key"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# --- Users Persistence ---
USERS_FILE = "users.json"
if not os.path.exists(USERS_FILE):
    initial_users = {
        "admin": {"username": "admin", "hashed_password": pwd_context.hash("adminpass"), "role": "Admin"}
    }
    with open(USERS_FILE, "w") as f:
        json.dump(initial_users, f, indent=4)

def load_users():
    try:
        with open(USERS_FILE, "r") as f:
            return json.load(f)
    except json.JSONDecodeError:
        logging.error(f"Users file '{USERS_FILE}' is corrupted. Reinitializing.")
        # Reinitialize if corrupted
        initial_users = {
            "admin": {"username": "admin", "hashed_password": pwd_context.hash("adminpass"), "role": "Admin"}
        }
        with open(USERS_FILE, "w") as f:
            json.dump(initial_users, f, indent=4)
        return initial_users
    except Exception as e:
        logging.error(f"Failed to load users file: {e}")
        return {}

def save_users(users):
    try:
        with open(USERS_FILE, "w") as f:
            json.dump(users, f, indent=4)
    except Exception as e:
        logging.error(f"Failed to save users file: {e}")

# --- Encryption Key Management ---
ENCRYPTION_KEY_PATH = "encryption.key"

def initialize_encryption_key():
    if not os.path.exists(ENCRYPTION_KEY_PATH):
        logging.info("No encryption key found. Generating new key.")
        key = Fernet.generate_key()
        try:
            with open(ENCRYPTION_KEY_PATH, "wb") as key_file:
                key_file.write(key)
            Fernet(key)  # Validate new key
            logging.info("New encryption key generated and saved.")
            return key
        except Exception as e:
            logging.error(f"Failed to save new encryption key: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to initialize encryption key: {e}")
    try:
        with open(ENCRYPTION_KEY_PATH, "rb") as key_file:
            key = key_file.read()
        if not key:
            logging.error("Encryption key file is empty. Generating new key.")
            key = Fernet.generate_key()
            with open(ENCRYPTION_KEY_PATH, "wb") as key_file:
                key_file.write(key)
            logging.warning("New encryption key generated due to empty key file.")
            return key
        # Validate key by attempting a test encryption/decryption
        fernet_test = Fernet(key)
        test_data = b"test"
        test_encrypted = fernet_test.encrypt(test_data)
        fernet_test.decrypt(test_encrypted)
        logging.info("Encryption key loaded and validated successfully.")
        return key
    except (InvalidToken, Exception) as e:
        logging.error(f"Invalid encryption key at {ENCRYPTION_KEY_PATH}: {str(e)}. Generating new key.")
        key = Fernet.generate_key()
        try:
            with open(ENCRYPTION_KEY_PATH, "wb") as key_file:
                key_file.write(key)
            Fernet(key)  # Validate new key
            logging.warning("New encryption key generated. Previous files may be inaccessible.")
            return key
        except Exception as e:
            logging.error(f"Failed to save new encryption key: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to initialize encryption key: {e}")

fernet = Fernet(initialize_encryption_key())

# --- Ensure Directories Exist ---
try:
    os.makedirs("encrypted_files", exist_ok=True)
    os.makedirs("logs", exist_ok=True)
    os.makedirs("user_memory", exist_ok=True)
    os.makedirs("tmp", exist_ok=True)  # Ensure 'tmp' directory exists for temporary files
except Exception as e:
    logging.error(f"Failed to create directories: {e}")
    raise HTTPException(status_code=500, detail=f"Failed to initialize directories: {e}")

# --- Logging Configuration ---
logging.basicConfig(filename="logs/audit.log", level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Embeddings & Vector DB Setup ---
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
client = chromadb.PersistentClient(path="./chroma_db")
collection = client.get_or_create_collection("rag_collection")

# --- Clean up invalid documents in vector store on startup ---
def clean_vector_store():
    try:
        documents = collection.get(include=["metadatas"])
        if not documents or not documents["ids"]:
            logging.info("Vector store is empty or no documents to clean.")
            return

        invalid_ids = []
        for i, meta in enumerate(documents["metadatas"]):
            file_id = meta.get("file_id")
            doc_id = documents["ids"][i]

            if not file_id:
                invalid_ids.append(doc_id)
                logging.warning(f"Document {doc_id} has no file_id in metadata. Marking for deletion.")
                continue

            enc_path = f"encrypted_files/{file_id}"
            meta_path = f"encrypted_files/{file_id}.meta"

            if not os.path.exists(enc_path):
                invalid_ids.append(doc_id)
                logging.warning(f"Encrypted file {file_id} missing. Marking document {doc_id} for deletion.")
                continue
            if not os.path.exists(meta_path):
                invalid_ids.append(doc_id)
                logging.warning(f"Metadata file {file_id}.meta missing. Marking document {doc_id} for deletion.")
                continue

            try:
                with open(enc_path, "rb") as ef:
                    enc_bytes = ef.read()
                if not enc_bytes:
                    invalid_ids.append(doc_id)
                    logging.warning(f"Encrypted file {file_id} is empty. Marking document {doc_id} for deletion.")
                    continue
                fernet.decrypt(enc_bytes)  # Full decryption for validation
            except Exception as e:
                invalid_ids.append(doc_id)
                logging.warning(f"File {file_id} failed decryption: {str(e)}. Marking document {doc_id} for deletion.")

        if invalid_ids:
            collection.delete(ids=invalid_ids)
            logging.info(f"Deleted {len(invalid_ids)} invalid documents from vector store: {invalid_ids}")
        else:
            logging.info("No invalid documents found in vector store.")
    except Exception as e:
        logging.error(f"Failed to clean vector store: {e}")

# Run cleanup on startup
clean_vector_store()

# --- LLM Setup ---
llm = OllamaLLM(model="phi")  # You might need to pull this model first (ollama pull phi)

# --- Pydantic Models ---
class Token(BaseModel):
    access_token: str
    token_type: str
    role: str

class TokenData(BaseModel):
    username: str | None = None

class User(BaseModel):
    username: str
    role: str

class NewUser(BaseModel):
    username: str
    password: str
    role: str = "User"

class Citation(BaseModel):
    text: str
    file_id: Optional[str] = None
    file_type: Optional[str] = None  # Added for frontend preview logic
    page_num: Optional[int] = None  # Added for PDF preview
    start_time: Optional[float] = None  # Added for Audio preview (seconds)
    original_filename: Optional[str] = None  # Added for preview display

class Message(BaseModel):
    sender: str
    text: str
    timestamp: str
    citations: Optional[List[Citation]] = None

class Conversation(BaseModel):
    id: str
    timestamp: str
    messages: List[Message]

class QueryRequest(BaseModel):
    question: str
    conversation_id: Optional[str] = None

class QueryResponse(BaseModel):
    answer: str
    citations: List[Citation]
    conversation_id: str

# --- Security Helpers ---
def verify_password(plain, hashed):
    return pwd_context.verify(plain, hashed)

def get_user(username):
    users = load_users()
    return users.get(username)

def authenticate_user(username, password):
    user = get_user(username)
    if not user or not verify_password(password, user["hashed_password"]):
        return False
    return user

def create_access_token(data: dict, expires_delta: timedelta | None = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=15))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(status_code=401, detail="Could not validate credentials")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        role: str = payload.get("role", "Guest")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = get_user(username)
    if not user:
        raise credentials_exception
    return User(username=user["username"], role=user["role"])

# --- Conversation Memory Manager ---
class ConversationManager:
    def __init__(self, username: str):
        self.username = username
        self.mem_path = f"user_memory/{self.username}.json"

    def _load_conversations_data(self) -> List[dict]:
        if not os.path.exists(self.mem_path):
            return []
        try:
            with open(self.mem_path, "r") as f:
                return json.load(f)
        except json.JSONDecodeError:
            logging.error(f"Memory for user {self.username} at '{self.mem_path}' is corrupted. Starting new memory.")
            return []
        except Exception as e:
            logging.error(f"Error loading memory for user {self.username}: {e}")
            return []

    def _save_conversations_data(self, conversations_data: List[dict]):
        try:
            with open(self.mem_path, "w") as f:
                json.dump(conversations_data, f, indent=4)
        except Exception as e:
            logging.error(f"Failed to save memory for user {self.username}: {e}")

    def get_conversations(self) -> List[Conversation]:
        raw_conversations = self._load_conversations_data()
        validated_conversations: List[Conversation] = []
        needs_save = False

        for conv_data in raw_conversations:
            conv_id = conv_data.get("id")
            if not conv_id or not isinstance(conv_id, str):
                conv_id = str(uuid.uuid4())
                conv_data["id"] = conv_id
                needs_save = True  # Mark for saving if an ID was generated

            messages: List[Message] = []
            for msg_data in conv_data.get("messages", []):
                citations: List[Citation] = []
                if msg_data.get("citations") is not None:
                    for cit_data in msg_data.get("citations", []):
                        file_id = cit_data.get("file_id")
                        if file_id:
                            enc_path = f"encrypted_files/{file_id}"
                            meta_path = f"encrypted_files/{file_id}.meta"
                            if os.path.exists(enc_path) and os.path.exists(meta_path):
                                try:
                                    with open(enc_path, "rb") as ef:
                                        enc_bytes = ef.read()
                                    if enc_bytes:
                                        fernet.decrypt(enc_bytes)  # Validate decryption
                                        citations.append(Citation(
                                            text=cit_data.get("text", ""),
                                            file_id=file_id,
                                            file_type=cit_data.get("file_type"),
                                            page_num=cit_data.get("page_num"),
                                            start_time=cit_data.get("start_time"),
                                            original_filename=cit_data.get("original_filename")
                                        ))
                                except (InvalidToken, Exception) as e:
                                    logging.warning(f"Skipping invalid file {file_id} during conversation load due to decryption failure: {e}")
                            else:
                                logging.warning(f"Skipping citation for missing file {file_id} during conversation load.")

                messages.append(Message(
                    sender=msg_data.get("sender", "unknown"),
                    text=msg_data.get("text", ""),
                    timestamp=msg_data.get("timestamp", datetime.utcnow().isoformat()),
                    citations=citations if citations else None
                ))
            validated_conversations.append(Conversation(
                id=conv_id,
                timestamp=conv_data.get("timestamp", datetime.utcnow().isoformat()),
                messages=messages
            ))

        if needs_save:  # Only save if changes were made (e.g., new UUIDs generated)
            self._save_conversations_data([c.dict() for c in validated_conversations])  # Save the validated structure

        validated_conversations.sort(key=lambda c: c.timestamp, reverse=True)  # Sort newest first
        return validated_conversations

    def add_message_turn(self, conversation_id: str, user_message: Message, ai_message: Message):
        conversations_data = self._load_conversations_data()

        target_conv = None
        for conv in conversations_data:
            if conv.get("id") == conversation_id:
                target_conv = conv
                break

        if target_conv:
            target_conv["messages"].append(user_message.dict())
            target_conv["messages"].append(ai_message.dict())
            target_conv["timestamp"] = ai_message.timestamp  # Update timestamp for sorting
        else:
            new_conversation = {
                "id": conversation_id,
                "timestamp": ai_message.timestamp,  # Timestamp of the first message turn
                "messages": [user_message.dict(), ai_message.dict()]
            }
            conversations_data.append(new_conversation)

        self._save_conversations_data(conversations_data)

# --- Auth Endpoints ---
@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    user = authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    token = create_access_token(data={"sub": user["username"], "role": user["role"]},
                                expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    logging.info(f"Login: {user['username']}")
    return {"access_token": token, "token_type": "bearer", "role": user["role"]}

# --- User Management Endpoints ---
@app.get("/users/me", response_model=User)
async def get_me(current_user: User = Depends(get_current_user)):
    return current_user

@app.get("/users", response_model=List[User])
async def list_users(current_user: User = Depends(get_current_user)):
    if current_user.role != "Admin":
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    users = load_users()
    return [User(username=u, role=users[u]["role"]) for u in users]

@app.post("/users", status_code=201)
async def create_user(new_user: NewUser, current_user: User = Depends(get_current_user)):
    if current_user.role != "Admin":
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    users = load_users()
    if new_user.username in users:
        raise HTTPException(status_code=400, detail="User with this username already exists")
    hashed_password = pwd_context.hash(new_user.password)
    users[new_user.username] = {"username": new_user.username, "hashed_password": hashed_password, "role": new_user.role}
    save_users(users)
    logging.info(f"User created: {new_user.username} by {current_user.username}")
    return {"message": "User created successfully"}

@app.delete("/users/{username}", status_code=204)
async def delete_user(username: str, current_user: User = Depends(get_current_user)):
    if current_user.role != "Admin":
        raise HTTPException(status_code=403, detail="Insufficient permissions")
    if username == "admin":
        raise HTTPException(status_code=400, detail="Cannot delete the default admin user")
    users = load_users()
    if username not in users:
        raise HTTPException(status_code=404, detail="User not found")
    del users[username]
    save_users(users)
    logging.info(f"User deleted: {username} by {current_user.username}")
    return {"message": "User deleted"}

# --- File Upload & Ingestion ---
@app.post("/upload")
async def upload_file(file: UploadFile = File(...), current_user: User = Depends(get_current_user)):
    if current_user.role not in ["Admin", "Analyst", "Guest", "User"]:
        raise HTTPException(status_code=403, detail="Insufficient permissions")

    content = await file.read()
    if not content or len(content) == 0:
        logging.error(f"Uploaded file {file.filename} is empty or unreadable.")
        raise HTTPException(status_code=400, detail="Uploaded file is empty or unreadable")

    file_id = str(uuid.uuid4())
    encrypted_path = f"encrypted_files/{file_id}"
    meta_path = f"encrypted_files/{file_id}.meta"

    try:
        # Test encryption with small data
        test_data = b"test_data"
        test_encrypted = fernet.encrypt(test_data)
        fernet.decrypt(test_encrypted)  # Ensure key is valid
        encrypted_content = fernet.encrypt(content)
        with open(encrypted_path, "wb") as f:
            f.write(encrypted_content)
        # Verify written file
        with open(encrypted_path, "rb") as ef:
            enc_bytes = ef.read()
        if not enc_bytes or len(enc_bytes) != len(encrypted_content):
            raise Exception("Written encrypted file is empty or corrupted")
        fernet.decrypt(enc_bytes)  # Full decryption to validate
    except InvalidToken as e:
        logging.error(f"Encryption/decryption failed for file {file.filename} (ID: {file_id}): Invalid token - {str(e)}")
        if os.path.exists(encrypted_path):
            os.remove(encrypted_path)
        raise HTTPException(status_code=500, detail="Failed to secure file storage: Invalid encryption key")
    except Exception as e:
        logging.error(f"Failed to encrypt or verify file {file.filename} (ID: {file_id}): {str(e)}")
        if os.path.exists(encrypted_path):
            os.remove(encrypted_path)
        raise HTTPException(status_code=500, detail=f"Failed to secure file storage: {str(e)}")

    file_meta = {
        "file_id": file_id,
        "original_filename": file.filename,
        "content_type": file.content_type,
        "uploader": current_user.username,
        "uploaded_at": datetime.utcnow().isoformat()
    }
    try:
        with open(meta_path, "w") as m:
            json.dump(file_meta, m, indent=4)
    except Exception as e:
        logging.error(f"Failed to save metadata for file {file_id}: {e}")
        if os.path.exists(encrypted_path):
            os.remove(encrypted_path)
        raise HTTPException(status_code=500, detail=f"Failed to save file metadata: {e}")

    base_metadata = {
        "file_id": file_id,
        "source": file.filename,
        "uploader": current_user.username
    }

    vectorstore = Chroma(client=client, collection_name="rag_collection", embedding_function=embeddings)
    text_content = ""
    file_type_inferred = "UNKNOWN"  # Default type

    temp_file_path = None  # Initialize to None

    try:
        # Create tempfile in the "tmp" directory
        with tempfile.NamedTemporaryFile(delete=False, dir="tmp", suffix=os.path.splitext(file.filename)[1]) as temp_file:
            temp_file.write(content)
            temp_file_path = temp_file.name

        try:
            if file.filename.lower().endswith(".pdf"):
                file_type_inferred = "PDF"
                try:
                    reader = PdfReader(temp_file_path)
                    page_texts = []
                    for i, page in enumerate(reader.pages):
                        ptext = page.extract_text()
                        if ptext and ptext.strip():
                            page_texts.append(ptext)
                            metadata = {**base_metadata, "file_type": "PDF", "page": i + 1}
                            doc = Document(page_content=ptext, metadata=metadata)
                            vectorstore.add_documents([doc])
                    text_content = " ".join(page_texts)
                except Exception as e:
                    logging.warning(f"Failed to process PDF {file.filename} (ID: {file_id}): {e}")
                    text_content = ""
            elif file.filename.lower().endswith(".docx"):
                file_type_inferred = "DOCX"
                try:
                    docx_doc = DocxDocument(temp_file_path)
                    text_content = " ".join(para.text for para in docx_doc.paragraphs if para.text and para.text.strip())
                    if text_content:
                        metadata = {**base_metadata, "file_type": "DOCX"}
                        doc = Document(page_content=text_content, metadata=metadata)
                        vectorstore.add_documents([doc])
                except Exception as e:
                    logging.warning(f"Failed to process DOCX {file.filename} (ID: {file_id}): {e}")
                    text_content = ""
            elif file.filename.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff")):  # More image types
                file_type_inferred = "IMAGE"
                try:
                    reader = easyocr.Reader(['en'])
                    ocr_text = reader.readtext(temp_file_path, detail=0)
                    text_content = " ".join(ocr_text) if ocr_text else ""
                    if text_content:
                        metadata = {**base_metadata, "file_type": "IMAGE"}
                        doc = Document(page_content=text_content, metadata=metadata)
                        vectorstore.add_documents([doc])
                except Exception as e:
                    logging.warning(f"Failed to process image {file.filename} (ID: {file_id}): {e}")
                    text_content = ""
            elif file.filename.lower().endswith((".mp3", ".wav", ".ogg", ".flac")):  # More audio types
                file_type_inferred = "AUDIO"
                try:
                    model = whisper.load_model("base")
                    result = model.transcribe(temp_file_path)
                    text_content = result.get("text", "")
                    # Extract start time of the first segment for audio citation
                    ts = None
                    segments = result.get("segments")
                    if segments and len(segments) > 0:
                        ts = segments[0].get("start")
                    if text_content:
                        metadata = {**base_metadata, "file_type": "AUDIO", "timestamp": ts}
                        doc = Document(page_content=text_content, metadata=metadata)
                        vectorstore.add_documents([doc])
                except Exception as e:
                    logging.warning(f"Failed to process audio {file.filename} (ID: {file_id}): {e}")
                    text_content = ""
            else:  # Treat as generic text file if no specific handler matches
                file_type_inferred = "TEXT"
                try:
                    text_content = content.decode(errors="ignore")
                    if text_content.strip():
                        metadata = {**base_metadata, "file_type": "TEXT"}
                        doc = Document(page_content=text_content, metadata=metadata)
                        vectorstore.add_documents([doc])
                except Exception as e:
                    logging.warning(f"Failed to process text file {file.filename} (ID: {file_id}): {e}")
                    text_content = ""

        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except Exception as e:
                    logging.warning(f"Failed to remove temp file {temp_file_path}: {e}")

        # Update metadata with the inferred file_type
        file_meta["file_type"] = file_type_inferred
        with open(meta_path, "w") as m:
            json.dump(file_meta, m, indent=4)

        if not text_content.strip():
            logging.warning(f"No valid text extracted from {file.filename} (ID: {file_id}). Skipping embedding.")
            # Only remove if the file was actually uploaded and failed ingestion
            if os.path.exists(encrypted_path):
                os.remove(encrypted_path)
            if os.path.exists(meta_path):
                os.remove(meta_path)
            raise HTTPException(status_code=400, detail=f"No valid content extracted from file: {file.filename}")
    except Exception as e:
        logging.error(f"Error during file ingestion for {file.filename} (ID: {file_id}): {e}")
        # Ensure cleanup if ingestion fails at any stage
        if os.path.exists(encrypted_path):
            os.remove(encrypted_path)
        if os.path.exists(meta_path):
            os.remove(meta_path)
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        raise HTTPException(status_code=500, detail=f"Failed to process and ingest file: {e}")

    logging.info(f"File uploaded and ingested: {file.filename} (ID: {file_id}) ({file_type_inferred}) by {current_user.username}")
    return {"message": "File uploaded and ingested", "file_id": file_id, "file_type": file_type_inferred}

# --- Query with Citations and Conversation Memory ---
@app.post("/query", response_model=QueryResponse)
async def query_rag(query_request: QueryRequest, current_user: User = Depends(get_current_user)):
    vectorstore = Chroma(client=client, collection_name="rag_collection", embedding_function=embeddings)

    # Configure retriever filters based on user role
    search_kwargs = {"k": 5}
    # Admins can see all documents; other roles only see their own uploaded documents
    if current_user.role != "Admin":
        search_kwargs["filter"] = {"uploader": current_user.username}

    docs_and_scores = []
    try:
        # Using .asimilarity_search_with_score for async operation
        docs_and_scores = await vectorstore.asimilarity_search_with_score(query_request.question, **search_kwargs)
        # Filter for documents with a score above a certain threshold if needed, or take top K
        relevant_docs = [doc for doc, score in docs_and_scores if score > 0.3]  # Example threshold
    except Exception as e:
        logging.warning(f"Vector store search failed for query: '{query_request.question}'. Error: {e}")
        relevant_docs = []

    context = ""
    citations_for_response: List[Citation] = []
    seen_file_ids = set()  # To avoid duplicate citations for the same file_id

    for i, doc in enumerate(relevant_docs, start=1):
        meta = doc.metadata or {}
        file_type = meta.get("file_type", "Unknown")
        source = meta.get("source", "Unknown")
        file_id = meta.get("file_id")

        if file_id and file_id not in seen_file_ids:
            enc_path = f"encrypted_files/{file_id}"
            meta_path = f"encrypted_files/{file_id}.meta"
            if not os.path.exists(enc_path):
                logging.warning(f"Encrypted file {file_id} missing for query context.")
                continue
            if not os.path.exists(meta_path):
                logging.warning(f"Metadata file {file_id}.meta missing for query context.")
                continue
            try:
                with open(enc_path, "rb") as ef:
                    enc_bytes = ef.read()
                if not enc_bytes:
                    logging.warning(f"Encrypted file {file_id} is empty.")
                    continue
                fernet.decrypt(enc_bytes)  # Attempt decryption to validate file integrity

                cit_text = f"[{i}] {source} ({file_type})"
                page_num = meta.get("page")
                start_time = meta.get("timestamp")  # For audio

                if file_type.upper() == "PDF" and page_num:
                    cit_text += f", page {page_num}"
                if file_type.upper() == "AUDIO" and start_time is not None:
                    cit_text += f", time {start_time:.1f}s"  # Display seconds for audio

                citations_for_response.append(Citation(
                    text=cit_text,
                    file_id=file_id,
                    file_type=file_type,
                    page_num=page_num,
                    start_time=start_time,
                    original_filename=source  # Use 'source' as original_filename
                ))
                seen_file_ids.add(file_id)
                context += doc.page_content + "\n\n"
            except InvalidToken:
                logging.warning(f"File {file_id} failed decryption: Invalid token. Possible key mismatch.")
                continue
            except Exception as e:
                logging.warning(f"File {file_id} failed decryption: {str(e)}.")
                continue

    if context:
        try:
            answer = llm.invoke(f"Answer the question based only on the following context, use citations in the format [citation_number] where appropriate, e.g., 'The main point is X [1].'\nContext:\n{context}\n\nQuestion: {query_request.question}")
            if not answer or not answer.strip():
                answer = "No relevant information found in the provided context."
        except Exception as e:
            logging.error(f"LLM failed with context for query '{query_request.question}': {e}")
            answer = "Sorry, I couldn't generate an answer from the context due to an internal error."
    else:
        try:
            answer = llm.invoke(f"Answer the question concisely: {query_request.question}")
            if not answer or not answer.strip():
                answer = "Sorry, I couldn't generate an answer. Please try rephrasing your question."
        except Exception as e:
            logging.error(f"LLM failed for query '{query_request.question}': {e}")
            answer = "Sorry, I couldn't generate an answer due to an internal error."

    conv_manager = ConversationManager(current_user.username)
    current_time = datetime.utcnow().isoformat()
    conversation_id = query_request.conversation_id if query_request.conversation_id else str(uuid.uuid4())
    user_message = Message(sender="user", text=query_request.question, timestamp=current_time)
    ai_message = Message(sender="ai", text=answer, timestamp=current_time, citations=citations_for_response if citations_for_response else None)
    conv_manager.add_message_turn(conversation_id, user_message, ai_message)

    return QueryResponse(answer=answer, citations=citations_for_response, conversation_id=conversation_id)

# --- Get user memory (conversation history) ---
@app.get("/memory", response_model=List[Conversation])
async def get_memory(current_user: User = Depends(get_current_user)):
    conv_manager = ConversationManager(current_user.username)
    conversations = conv_manager.get_conversations()
    return conversations

# --- View (download/serve) uploaded file ---
@app.get("/view/{file_id}")
async def view_file(file_id: str, current_user: User = Depends(get_current_user)):
    meta_path = f"encrypted_files/{file_id}.meta"
    enc_path = f"encrypted_files/{file_id}"

    if not os.path.exists(meta_path) or not os.path.exists(enc_path):
        raise HTTPException(status_code=404, detail="File or metadata not found. It may have been deleted or not uploaded correctly.")

    try:
        with open(meta_path, "r") as m:
            meta = json.load(m)
    except Exception as e:
        logging.error(f"Failed to read meta file {meta_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to read file metadata: {e}")

    if current_user.role != "Admin" and meta.get("uploader") != current_user.username:
        raise HTTPException(status_code=403, detail="Insufficient permissions to view this file")

    try:
        with open(enc_path, "rb") as ef:
            enc_bytes = ef.read()
        if not enc_bytes:
            raise HTTPException(status_code=500, detail="Encrypted file is empty or corrupted.")
    except Exception as e:
        logging.error(f"Failed to read encrypted file {enc_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to read encrypted file: {e}")

    try:
        decrypted = fernet.decrypt(enc_bytes)
    except InvalidToken:
        logging.error(f"Decryption failed for file_id {file_id}: Invalid token.")
        raise HTTPException(status_code=500, detail=f"Failed to decrypt file: Invalid or corrupt encryption data. Key might have changed.")
    except Exception as e:
        logging.error(f"Decryption failed for file_id {file_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to decrypt file: {e}")

    original_filename = meta.get("original_filename", f"{file_id}")
    _, file_extension = os.path.splitext(original_filename)

    try:
        # Create tempfile in the "tmp" directory
        tmp_file = tempfile.NamedTemporaryFile(delete=False, dir="tmp", suffix=file_extension)
        tmp_path = tmp_file.name
        tmp_file.write(decrypted)
        tmp_file.close()
    except Exception as e:
        logging.error(f"Failed to write temp file {tmp_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to prepare file for viewing: {e}")

    return FileResponse(
        tmp_path,
        filename=original_filename,
        media_type=meta.get("content_type", "application/octet-stream"),
        background=asyncio.create_task(task_cleanup_temp_file(tmp_path))
    )

# Cleanup temp file in background (for /view endpoint)
async def task_cleanup_temp_file(tmp_path: str):
    try:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        logging.info(f"Cleaned up temporary file: {tmp_path}")
    except Exception as e:
        logging.error(f"Failed to clean up temp file {tmp_path}: {e}")

# --- New /preview endpoint for in-browser content display ---
@app.get("/preview/{file_id}")
async def preview_file(
    file_id: str,
    current_user: User = Depends(get_current_user),
    page_num: Optional[int] = None,
    start_time: Optional[float] = None
):
    meta_path = f"encrypted_files/{file_id}.meta"
    enc_path = f"encrypted_files/{file_id}"

    if not os.path.exists(meta_path) or not os.path.exists(enc_path):
        logging.error(f"Preview: File {file_id} or metadata not found.")
        raise HTTPException(status_code=404, detail="File or metadata not found for preview.")

    try:
        with open(meta_path, "r") as m:
            meta = json.load(m)
    except Exception as e:
        logging.error(f"Preview: Failed to read meta file {meta_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to read file metadata: {e}")

    if current_user.role != "Admin" and meta.get("uploader") != current_user.username:
        logging.warning(f"Preview: User {current_user.username} unauthorized for file {file_id}.")
        raise HTTPException(status_code=403, detail="Insufficient permissions to preview this file")

    try:
        with open(enc_path, "rb") as ef:
            enc_bytes = ef.read()
        if not enc_bytes:
            logging.error(f"Preview: Encrypted file {file_id} is empty.")
            raise HTTPException(status_code=500, detail="Encrypted file is empty or corrupted.")
    except Exception as e:
        logging.error(f"Preview: Failed to read encrypted file {enc_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to read encrypted file: {e}")

    try:
        decrypted_content = fernet.decrypt(enc_bytes)
    except InvalidToken:
        logging.error(f"Preview: Decryption failed for file_id {file_id}: Invalid token.")
        raise HTTPException(status_code=500, detail="Failed to decrypt file: Invalid or corrupt encryption data.")
    except Exception as e:
        logging.error(f"Preview: Decryption failed for file_id {file_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to decrypt file: {e}")

    file_type = meta.get("file_type", "UNKNOWN").upper()
    content_type = meta.get("content_type", "application/octet-stream")

    if file_type == "PDF":
        if page_num is None:
            logging.warning(f"Preview: Page number missing for PDF {file_id}.")
            raise HTTPException(status_code=400, detail="Page number is required for PDF preview.")
        try:
            pdf_document = fitz.open(stream=decrypted_content, filetype="pdf")
            if 0 <= page_num - 1 < len(pdf_document):
                page = pdf_document.load_page(page_num - 1)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                img_bytes = pix.tobytes("png")
                pdf_document.close()
                logging.info(f"Preview: Served PDF page {page_num} for {file_id} as PNG.")
                return StreamingResponse(BytesIO(img_bytes), media_type="image/png")
            else:
                pdf_document.close()
                logging.warning(f"Preview: Page {page_num} not found in PDF {file_id}.")
                raise HTTPException(status_code=404, detail=f"Page {page_num} not found in PDF.")
        except Exception as e:
            logging.error(f"Preview: Failed to render PDF page {page_num} for {file_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to render PDF page: {e}")

    elif file_type == "IMAGE":
        logging.info(f"Preview: Served image {file_id} directly.")
        return StreamingResponse(BytesIO(decrypted_content), media_type=content_type)

    elif file_type == "AUDIO":
        logging.info(f"Preview: Served audio {file_id} directly.")
        return StreamingResponse(BytesIO(decrypted_content), media_type=content_type)

    elif file_type == "DOCX":
        try:
            with tempfile.NamedTemporaryFile(delete=False, dir="tmp", suffix=".docx") as tmp_docx:
                tmp_docx.write(decrypted_content)
                tmp_docx_path = tmp_docx.name
            doc = DocxDocument(tmp_docx_path)
            full_text = "\n".join(para.text for para in doc.paragraphs)
            os.unlink(tmp_docx_path)
            logging.info(f"Preview: Served DOCX {file_id} as plain text.")
            return StreamingResponse(BytesIO(full_text.encode('utf-8')), media_type="text/plain")
        except Exception as e:
            logging.error(f"Preview: Failed to process DOCX for {file_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to process DOCX: {e}")

    elif file_type == "TEXT":
        try:
            text_content = decrypted_content.decode('utf-8', errors='ignore')
            logging.info(f"Preview: Served TEXT file {file_id} directly.")
            return StreamingResponse(BytesIO(text_content.encode('utf-8')), media_type="text/plain")
        except Exception as e:
            logging.error(f"Preview: Failed to decode TEXT file {file_id}: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to decode text file: {e}")

    else:
        logging.warning(f"Preview: Unsupported file type {file_type} for {file_id}.")
        raise HTTPException(status_code=400, detail=f"Preview not available for file type: {file_type}")
if __name__ == "__main__":
    import uvicorn
    # Make sure to run with --reload if you are developing, remove for production
    uvicorn.run(app, host="0.0.0.0", port=8000)